from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Consolas'
style.font.size = Pt(9)
style.font.color.rgb = RGBColor(0, 0, 0)

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    h_style.font.bold = True
    h_style.font.name = 'Consolas'
    h_style.font.size = Pt(14) if level == 1 else Pt(12) if level == 2 else Pt(11)

def add_code(code, size=8):
    for line in code.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10)

doc.add_heading('PASSWORD AUDIT SYSTEM', level=1)

doc.add_heading('password_audit.py', level=2)
with open('password_audit.py', 'r', encoding='utf-8') as f:
    add_code(f.read(), size=7)

doc.add_page_break()
doc.add_heading('test_password_audit.py', level=2)
with open('test_password_audit.py', 'r', encoding='utf-8') as f:
    add_code(f.read(), size=7)

doc.add_page_break()
doc.add_heading('passwords.txt', level=2)
with open('passwords.txt', 'r', encoding='utf-8') as f:
    add_code(f.read(), size=9)

output_path = 'Password_Audit_System_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
